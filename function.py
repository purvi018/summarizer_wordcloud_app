import argparse
from gensim.summarization import summarize
from docx import Document

def get_summary(text,output_file_name):
    gensim_summary_text = summarize(text, word_count=200, ratio = 0.1)
    # Create a Word document
    doc = Document()
    doc.add_paragraph(gensim_summary_text)

    # Save the document as a .docx file
    doc.save(output_file_name)


def read_text_from_file(file_path):
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)




