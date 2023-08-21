from wordcloud import WordCloud
from docx import Document
import matplotlib.pyplot as plt
import io

def wordcloud_generator(text,output_file_name):
    # Generate the word cloud image
    wordcloud = WordCloud(width=300, height=300, background_color='white').generate(text)

    # Create a BytesIO stream to capture the image
    image_stream = io.BytesIO()
    wordcloud.to_image().save(image_stream, format='PNG')

    # Create a Word document
    doc = Document()
    doc.add_paragraph('Word Cloud Example')

    # Add the word cloud image to the document
    doc.add_picture(image_stream)

    # Save the document as a .docx file
    doc.save(output_file_name)


